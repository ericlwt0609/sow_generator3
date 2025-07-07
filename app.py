import streamlit as st
import openai
import tempfile
import os
from PyPDF2 import PdfReader
from docx import Document
from docx import Document as DocxWriter
import requests
from bs4 import BeautifulSoup

# Extract PDF text
def extract_pdf_text(file):
    reader = PdfReader(file)
    return "".join(page.extract_text() or "" for page in reader.pages)[:8000]

# Extract DOCX text
def extract_docx_text(file):
    doc = Document(file)
    return "\n".join(para.text for para in doc.paragraphs)[:8000]

# Fetch example SoW clauses from LawInsider
def fetch_lawinsider_examples():
    url = "https://www.lawinsider.com/clause/scope-of-work"
    response = requests.get(url)
    soup = BeautifulSoup(response.text, "html.parser")
    clauses = soup.select(".clause-body")
    return [clause.get_text(strip=True) for clause in clauses[:5]]

# Scrape clauses from a custom URL
def fetch_text_from_url(url):
    try:
        response = requests.get(url)
        soup = BeautifulSoup(response.text, "html.parser")
        paragraphs = soup.find_all('p')
        return "\n".join(p.get_text(strip=True) for p in paragraphs[:10])
    except Exception as e:
        return f"[Error fetching URL content: {e}]"

# Generate SoW using OpenAI
def generate_sow(base_text, user_desc, selected_examples, existing_sow=None, feedback=None):
    examples_text = "\n---\n".join(selected_examples) if selected_examples else "None included"

    if existing_sow and feedback:
        # If refining an existing SoW
        prompt = f'''
You are a legal AI assistant. Based on the existing Statement of Work and the user's feedback, generate an improved version of the SoW. Retain all parts that have not been highlighted or are not relevant to the changes requested in the feedback.

---
Existing Statement of Work:
{existing_sow}

---
User Feedback for Refinement:
{feedback}

---
Generate the improved SoW using the original structure:
1. Description – What is being supplied or done.
2. Function – The business purpose or outcome the goods/services serve.
3. Price – Pricing structure, billing frequency, and payment terms.
4. Milestones – Key deliverables with corresponding deadlines or phases.
5. Warranties – Any performance guarantees, service warranties, or coverage periods.
6. Service Levels (if applicable) – SLAs, KPIs, uptime, penalties, or escalation paths.
7. Others – Any additional relevant clauses not captured above (e.g. assumptions, subcontracting, ownership of deliverables).
'''
    else:
        # Initial generation
        prompt = f'''
You are a legal AI assistant. Based on the following base contract text, user description, and example SoW clauses, generate a detailed Statement of Work (SoW). You should also search the Internet for any information or SoW which are the same or similar as the subject in question and incorporate relevant information from there.

---
User Description:
{user_desc}

---
Base Document Extract:
{base_text}

---
Example SoWs:
{examples_text}

---
Generate the SoW using the following structure:

1. Description – What is being supplied or done.
2. Function – The business purpose or outcome the goods/services serve.
3. Price – Pricing structure, billing frequency, and payment terms.
4. Milestones – Key deliverables with corresponding deadlines or phases.
5. Warranties – Any performance guarantees, service warranties, or coverage periods.
6. Service Levels (if applicable) – SLAs, KPIs, uptime, penalties, or escalation paths.
7. Others – Any additional relevant clauses not captured above (e.g. assumptions, subcontracting, ownership of deliverables).

Also suggest questions for missing or unclear details.
'''

    client = openai.OpenAI(base_url="https://generativelanguage.googleapis.com/v1beta/openai/", api_key=os.getenv("OPENAI_API_KEY"))

    response = client.chat.completions.create(
        model="gemini-2.5-flash",
        messages=[
            {"role": "system", "content": "You are a contract lawyer."},
            {"role": "user", "content": prompt}
        ],
        temperature=0.5
    )
    return response.choices[0].message.content

# Export output to DOCX
def export_to_docx(content):
    doc = DocxWriter()
    doc.add_heading("Generated Statement of Work", level=1)
    for paragraph in content.split('\n'):
        doc.add_paragraph(paragraph)
    temp_path = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(temp_path.name)
    return temp_path.name

# Streamlit UI
st.title("AI Statement of Work (SoW) Generator")

uploaded_file = st.file_uploader("Upload all relevant/to-date client presentation, proposal, scope document and even base contracts which you wish this to be based on (PDF or DOCX)", type=["pdf", "docx"])
user_desc = st.text_area("Describe the goods/services and business context")

# LawInsider examples are now always included
with st.spinner("Fetching LawInsider examples..."):
    lawinsider_examples = fetch_lawinsider_examples()

custom_examples_input = st.text_area("Paste your own SoW clauses or content here (optional)")
external_url = st.text_input("Paste a URL to extract external SoW-style clauses (optional)")

if st.button("Generate SoW"):
    if uploaded_file and user_desc:
        with st.spinner("Extracting text and generating SoW..."):
            # Extract uploaded content
            if uploaded_file.name.endswith(".pdf"):
                text = extract_pdf_text(uploaded_file)
            else:
                text = extract_docx_text(uploaded_file)

            # Collect all examples, including LawInsider automatically
            combined_examples = lawinsider_examples.copy()
            if custom_examples_input.strip():
                combined_examples.append(custom_examples_input.strip())
            if external_url.strip():
                combined_examples.append(fetch_text_from_url(external_url.strip()))

            # Generate SoW
            sow = generate_sow(text, user_desc, combined_examples)

            # Store the generated SoW in session state for refinement
            st.session_state.generated_sow = sow
            st.session_state.base_text = text
            st.session_state.user_desc = user_desc
            st.session_state.combined_examples = combined_examples

            # Show result
            st.subheader("Generated Statement of Work")
            st.write(sow)

            # Offer download
            docx_path = export_to_docx(sow)
            with open(docx_path, "rb") as f:
                st.download_button("Download SoW as DOCX", f, file_name="Statement_of_Work.docx")
    else:
        st.warning("Please upload a document and provide a description.")

# Feedback section for refinement
if 'generated_sow' in st.session_state and st.session_state.generated_sow:
    st.markdown("---")
    st.subheader("Refine Statement of Work")
    feedback_input = st.text_area("What would you like to add, remove, or edit in the SoW?", key="feedback_sow")

    if st.button("Refine SoW"):
        if feedback_input.strip():
            with st.spinner("Refining Statement of Work..."):
                refined_sow = generate_sow(
                    st.session_state.base_text,
                    st.session_state.user_desc,
                    st.session_state.combined_examples, # Pass existing examples for context
                    existing_sow=st.session_state.generated_sow,
                    feedback=feedback_input
                )
                st.session_state.generated_sow = refined_sow # Update the stored SoW
                st.subheader("Refined Statement of Work")
                st.write(refined_sow)

                # Offer download for refined SoW
                docx_path = export_to_docx(refined_sow)
                with open(docx_path, "rb") as f:
                    st.download_button("Download Refined SoW as DOCX", f, file_name="Refined_Statement_of_Work.docx")
        else:
            st.warning("Please provide feedback to refine the SoW.")
